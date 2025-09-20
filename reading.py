"""
Reading and parsing of documents (.docx, .rtf).

- Uses `python-docx` for .docx files
- Uses `striprtf` for .rtf files
- Provides a single-threaded recursive reader that returns plain text and metadata
"""
from __future__ import annotations

import hashlib
import logging
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:  # pragma: no cover
    DocxDocument = None  # type: ignore

try:
    from striprtf.striprtf import rtf_to_text  # striprtf
except Exception:  # pragma: no cover
    rtf_to_text = None  # type: ignore


@dataclass
class ParsedDocument:
    """Container for parsed document content and metadata.

    Attributes:
        source_path: Full path to the source file.
        filename: Basename of the source file.
        text: Extracted plain text content.
        total_chars: Total character length of the extracted text.
        doc_hash: Optional hash of the text content for change detection.
    """
    source_path: str
    filename: str
    text: str
    total_chars: int
    doc_hash: str


class DocumentReader:
    """Reads and parses .docx and .rtf files under a directory tree."""

    SUPPORTED_EXTENSIONS = {".docx", ".rtf"}

    def __init__(self, logger: Optional[logging.Logger] = None) -> None:
        self.logger = logger or logging.getLogger(__name__)

    def read_documents(self, input_dir: str, file_glob: str = "**/*.{docx,rtf}", max_files: Optional[int] = None) -> List[ParsedDocument]:
        """Recursively read and parse documents.

        Args:
            input_dir: Root directory to scan.
            file_glob: Glob pattern to match files. Defaults to '**/*.{docx,rtf}'.
            max_files: Optional limit on number of files to process.

        Returns:
            A list of ParsedDocument objects.
        """
        base = Path(input_dir)
        if not base.exists() or not base.is_dir():
            raise ValueError(f"Input directory does not exist or is not a directory: {input_dir}")

        # pathlib does NOT support brace expansion like **/*.{docx,rtf}
        # Expand manually into multiple patterns if braces are present; otherwise, use as-is.
        patterns: List[str]
        if "{" in file_glob and "}" in file_glob:
            brace_content = file_glob[file_glob.find("{") + 1 : file_glob.find("}")]
            exts = [p.strip().lstrip(".") for p in brace_content.split(",") if p.strip()]
            prefix = file_glob[: file_glob.find("{")]
            suffix = file_glob[file_glob.find("}") + 1 :]
            patterns = [f"{prefix}{ext}{suffix}" for ext in exts]
        else:
            patterns = [file_glob]

        matched_paths: List[Path] = []
        for pat in patterns:
            # Use rglob if pattern starts with '**/' to ensure recursive search
            if pat.startswith("**/"):
                matched_paths.extend(base.rglob(pat[3:]))
            else:
                matched_paths.extend(base.glob(pat))

        # De-duplicate and filter by supported extensions
        seen = set()
        matches: List[Path] = []
        for p in matched_paths:
            if not p.is_file():
                continue
            ext = p.suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                continue
            rp = str(p.resolve())
            if rp in seen:
                continue
            seen.add(rp)
            matches.append(p)

        self.logger.info(f"Discovered {len(matches)} candidate files in {input_dir} with pattern(s) {patterns}")
        docs: List[ParsedDocument] = []
        for idx, path in enumerate(matches):
            if max_files is not None and idx >= max_files:
                break
            if not path.is_file():
                continue
            ext = path.suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                continue
            try:
                text = self._parse_file(path)
                text = text if text is not None else ""
                doc_hash = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
                pd = ParsedDocument(
                    source_path=str(path.resolve()),
                    filename=path.name,
                    text=text,
                    total_chars=len(text),
                    doc_hash=doc_hash,
                )
                docs.append(pd)
                self.logger.info(f"Parsed file: {path} (chars={len(text)})")
            except Exception as exc:  # pragma: no cover
                self.logger.warning(f"Failed to parse {path}: {exc}")
                continue
        return docs

    def _parse_file(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".docx":
            return self._parse_docx(path)
        if ext == ".rtf":
            return self._parse_rtf(path)
        raise ValueError(f"Unsupported extension: {ext}")

    def _parse_docx(self, path: Path) -> str:
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed")
        doc = DocxDocument(str(path))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        return "\n".join(parts)

    def _parse_rtf(self, path: Path) -> str:
        if rtf_to_text is None:
            raise RuntimeError("striprtf is not installed")
        data = Path(path).read_text(encoding="utf-8", errors="ignore")
        return rtf_to_text(data) or ""
